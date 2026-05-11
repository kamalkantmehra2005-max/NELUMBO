import streamlit as st
import pdfplumber
import re
from docx import Document
from datetime import datetime
from io import BytesIO


# =========================================================
# PAGE CONFIG
# =========================================================

st.set_page_config(
    page_title="Nelumbo",
    page_icon="🌸",
    layout="wide"
)


# =========================================================
# CUSTOM CSS
# =========================================================

st.markdown("""
<style>

.stApp {
    background-color: #0f1117;
    color: white;
}

h1,h2,h3,h4,p,div,label {
    color:white !important;
}

.stButton>button {
    background:#7c4dff;
    color:white;
    border:none;
    border-radius:10px;
    padding:12px;
    width:100%;
    font-size:16px;
    font-weight:bold;
}

.stDownloadButton>button {
    background:#00c853;
    color:white;
    border:none;
    border-radius:10px;
    padding:12px;
    width:100%;
    font-size:16px;
    font-weight:bold;
}

</style>
""", unsafe_allow_html=True)


# =========================================================
# HEADER
# =========================================================

st.markdown("""
<div style='text-align:center;'>

<div style='font-size:80px;'>
🌸
</div>

<h1>
Nelumbo
</h1>

<p>
Patent Automation Tool
</p>

</div>
""", unsafe_allow_html=True)


# =========================================================
# COUNTRY CODE MAP
# =========================================================

COUNTRY_CODES = {

    "CN": "Pepole's Republic of China",
    "JP": "Japan",
    "KR": "Republic of Korea",
    "US": "USA",
    "IN": "India",
    "EP": "European Patent Office",
    "DE": "Geremany"

}


# =========================================================
# ADDRESS SPLITTER
# =========================================================

def split_address(addr):

    addr = re.sub(r'\([A-Z]{2}\)', '', addr)

    parts = [p.strip() for p in addr.split(",")]

    house = parts[0] if len(parts) > 0 else ""

    street = parts[1] if len(parts) > 1 else ""

    city = parts[2] if len(parts) > 2 else ""

    state = parts[3] if len(parts) > 3 else ""

    country = parts[4] if len(parts) > 4 else ""

    pin = ""

    pin_match = re.search(r'(\d{6})', addr)

    if pin_match:

        pin = pin_match.group(1)

        state = state.replace(pin, "").strip()

    country = COUNTRY_CODES.get(
        country.upper(),
        country
    )

    return {

        "house": house,
        "street": street,
        "city": city,
        "state": state,
        "country": country,
        "pin": pin

    }


# =========================================================
# NORMALIZE TAGS
# =========================================================

def normalize_tag(tag):

    return re.sub(
        r'[^a-z0-9]',
        '',
        tag.lower()
    )


# =========================================================
# EXTRACT PDF DATA
# =========================================================

def extract_data(pdf_file):

    data = {}

    with pdfplumber.open(pdf_file) as pdf:

        text = " ".join([
            page.extract_text() or ""
            for page in pdf.pages
        ])

    text = re.sub(r'\s+', ' ', text)

    # =====================================================
    # COUNTRY CODE MAP
    # =====================================================

    country_map = {

        "CN": "China",
        "JP": "Japan",
        "KR": "Korea",
        "US": "United States",
        "IN": "India",
        "EP": "European Patent Office",
        "WO": "WIPO"

    }

    # =====================================================
    # APPLICANT
    # =====================================================

    applicant_match = re.search(
        r'Applicant\(s\):(.*?);(.*?)(?:\([A-Z]{2}\))',
        text
    )

    applicant_name = ""
    applicant_address = ""

    if applicant_match:

        applicant_name = applicant_match.group(1).strip()

        # REMOVE [CN/CN]
        applicant_name = re.sub(
            r'[\[\(]?[A-Z]{2}/[A-Z]{2}[\]\)]?',
            '',
            applicant_name
        ).strip()

        applicant_address = applicant_match.group(2).strip()

    data["applicant"] = applicant_name
    data["applicant_name"] = applicant_name

    # =====================================================
    # APPLICANT COUNTRY
    # =====================================================

    applicant_country_code = re.search(
        r'\(([A-Z]{2})\)',
        applicant_address
    )

    if applicant_country_code:

        data["country"] = country_map.get(
            applicant_country_code.group(1),
            applicant_country_code.group(1)
        )

    else:

        data["country"] = ""

    # =====================================================
    # ADDRESS SPLIT
    # =====================================================

    applicant_address = re.sub(
        r'\([A-Z]{2}\)',
        '',
        applicant_address
    )

    pin_match = re.search(
        r'(\d{6})',
        applicant_address
    )

    pin = pin_match.group(1) if pin_match else ""

    parts = [
        p.strip()
        for p in applicant_address.split(",")
    ]

    data["house"] = parts[0] if len(parts) > 0 else ""
    data["street"] = parts[1] if len(parts) > 1 else ""
    data["city"] = parts[2] if len(parts) > 2 else ""

    state = parts[3] if len(parts) > 3 else ""

    state = state.replace(pin, "").strip()

    data["state"] = state

    data["pin"] = pin

    # =====================================================
    # TITLE
    # =====================================================

    title_match = re.search(
        r'Title(?: of invention)?\s*:\s*(.*?)\s*(?:Abstract|Publication|PCT|Priority)',
        text,
        re.IGNORECASE
    )

    data["title"] = (
        title_match.group(1).strip()
        if title_match else ""
    )

    # =====================================================
    # PUBLICATION DATE
    # =====================================================

    pub_match = re.search(
        r'Publication date\s*:\s*(\d{2}\.\d{2}\.\d{4})',
        text,
        re.IGNORECASE
    )

    data["publication_date"] = (
        pub_match.group(1)
        if pub_match else ""
    )

    # =====================================================
    # PCT NUMBER
    # =====================================================

    pct_match = re.search(
        r'PCT/[A-Z]{2}\d{4}/\d+',
        text
    )

    data["application_no"] = (
        pct_match.group(0)
        if pct_match else ""
    )

    # =====================================================
    # PRIORITY
    # =====================================================

    priority_match = re.search(
        r'(\d{12}\.\w)',
        text
    )

    data["priority_no"] = (
        priority_match.group(1)
        if priority_match else ""
    )

    data["priority_country"] = "China"

    # =====================================================
    # FILING DATE
    # =====================================================

    filing_match = re.search(
        r'International filing date\s*:\s*(\d{2}\.\d{2}\.\d{4})',
        text,
        re.IGNORECASE
    )

    data["filing_date"] = (
        filing_match.group(1)
        if filing_match else ""
    )

    # =====================================================
    # INVENTOR SECTION
    # =====================================================

    inventor_section = re.search(
        r'\(72\)\s*Inventor\(s\):(.*?)\(74\)\s*Agent\(s\):',
        text,
        re.DOTALL | re.IGNORECASE
    )

    inventor_text = (
        inventor_section.group(1)
        if inventor_section else ""
    )

    inventor_pattern = re.findall(
        r'([A-Z][A-Z\s,\-]+);(.*?)(?=[A-Z][A-Z\s,\-]+;|$)',
        inventor_text
    )

    data["inventors"] = []

    inventor_names = []

    inventor_country = ""

    for idx, (name, address) in enumerate(
        inventor_pattern,
        start=1
    ):

        clean_name = name.strip()

        clean_name = re.sub(
            r'[\[\(]?[A-Z]{2}/[A-Z]{2}[\]\)]?',
            '',
            clean_name
        ).strip()

        inventor_names.append(clean_name)

        # COUNTRY
        country_match = re.search(
            r'\(([A-Z]{2})\)',
            address
        )

        country_name = ""

        if country_match:

            country_name = country_map.get(
                country_match.group(1),
                country_match.group(1)
            )

        inventor_country = country_name

        address = re.sub(
            r'\([A-Z]{2}\)',
            '',
            address
        )

        pin_match = re.search(
            r'(\d{6})',
            address
        )

        pin = pin_match.group(1) if pin_match else ""

        parts = [
            p.strip()
            for p in address.split(",")
        ]

        house = parts[0] if len(parts) > 0 else ""

        street = parts[1] if len(parts) > 1 else ""

        city = parts[2] if len(parts) > 2 else ""

        state = parts[3] if len(parts) > 3 else ""

        state = state.replace(pin, "").strip()

        inventor_data = {

            "name": clean_name,
            "house": house,
            "street": street,
            "city": city,
            "state": state,
            "country": country_name,
            "pin": pin

        }

        data["inventors"].append(
            inventor_data
        )

        data[f"inventor_{idx}_name"] = clean_name

        data[f"inventor_{idx}_country"] = country_name

    data["inventor_names"] = ", ".join(
        inventor_names
    )

    # =====================================================
    # MAIN COUNTRY TAG
    # =====================================================

    if inventor_country:
        data["country"] = inventor_country

    # =====================================================
    # TODAY DATES
    # =====================================================

    today = datetime.today()

    day = today.strftime("%d")

    month = today.strftime("%B")

    year = today.strftime("%Y")

    data["today_date_short"] = today.strftime(
        "%B %d, %Y"
    )

    data["today_date_long"] = (
        f"{day}th day of {month}, {year}"
    )

    return data

# =========================================================
# SMART TAG REPLACER
# =========================================================

def replace_in_runs(paragraph, data):

    for run in paragraph.runs:

        original_text = run.text

        tags = re.findall(
            r'{{(.*?)}}',
            original_text
        )

        updated_text = original_text

        for tag in tags:

            normalized_template_tag = normalize_tag(tag)

            replacement = ""

            for key, value in data.items():

                normalized_data_key = normalize_tag(key)

                if normalized_template_tag == normalized_data_key:

                    replacement = str(value)

                    break

            updated_text = updated_text.replace(
                "{{" + tag + "}}",
                replacement
            )

        run.text = updated_text


# =========================================================
# DOCUMENT GENERATOR
# =========================================================

def generate_doc(template_file, data):

    doc = Document(template_file)

    # =====================================================
    # PARAGRAPHS
    # =====================================================

    for para in doc.paragraphs:

        replace_in_runs(
            para,
            data
        )

    # =====================================================
    # TABLES
    # =====================================================

    for table in doc.tables:

        for row in table.rows:

            for cell in row.cells:

                for para in cell.paragraphs:

                    replace_in_runs(
                        para,
                        data
                    )

    output = BytesIO()

    doc.save(output)

    output.seek(0)

    return output


# =========================================================
# FILE UPLOADS
# =========================================================

st.markdown("## Upload Files")

pdf_file = st.file_uploader(
    "Upload PDF Information Sheet",
    type=["pdf"]
)

template_file = st.file_uploader(
    "Upload FORM-1 Template",
    type=["docx"]
)


# =========================================================
# GENERATE
# =========================================================

def generate_doc(template_file, data):

    doc = Document(template_file)

    # =====================================================
    # SMART TAG REPLACEMENT
    # =====================================================

    for para in doc.paragraphs:

        replace_in_runs(
            para,
            data
        )

    # =====================================================
    # TABLES
    # =====================================================

    for table in doc.tables:

        for row in table.rows:

            for cell in row.cells:

                for para in cell.paragraphs:

                    replace_in_runs(
                        para,
                        data
                    )

    # =====================================================
    # AUTO INVENTOR ROW CLONE
    # =====================================================

    for table in doc.tables:

        header_found = False

        for row in table.rows:

            row_text = " ".join([
                cell.text.lower()
                for cell in row.cells
            ])

            if "name in full" in row_text:

                header_found = True
                break

        if header_found:

            # KEEP HEADER ONLY
            while len(table.rows) > 1:

                tbl = table._tbl

                tbl.remove(
                    table.rows[-1]._tr
                )

            # ADD INVENTORS
            for inventor in data["inventors"]:

                # MAIN ROW
                row_cells = table.add_row().cells

                row_cells[0].text = inventor["name"]

                row_cells[1].text = "Unknown"

                row_cells[2].text = inventor["country"]

                # ADDRESS TITLE
                title_row = table.add_row().cells

                title_row[0].text = (
                    "Address of the Inventor"
                )

                title_row[0].merge(title_row[1])
                title_row[0].merge(title_row[2])

                # ADDRESS TABLE
                address_row = table.add_row().cells

                nested = address_row[0].add_table(
                    rows=6,
                    cols=2
                )

                nested.style = "Table Grid"

                labels = [

                    "House",
                    "Street",
                    "City",
                    "State",
                    "Country",
                    "Pin Code"

                ]

                values = [

                    inventor["house"],
                    inventor["street"],
                    inventor["city"],
                    inventor["state"],
                    inventor["country"],
                    inventor["pin"]

                ]

                for i in range(6):

                    nested.cell(i, 0).text = labels[i]

                    nested.cell(i, 1).text = values[i]

                address_row[0].merge(
                    address_row[1]
                )

                address_row[0].merge(
                    address_row[2]
                )

            break

    output = BytesIO()

    doc.save(output)

    output.seek(0)

    return output

# =========================================================
# FOOTER
# =========================================================

st.markdown("""
<hr>

<div style='text-align:center;color:gray;'>

Design by Kamal Kant<br>
Not recommended for convention and ordinary file

</div>
""", unsafe_allow_html=True)
